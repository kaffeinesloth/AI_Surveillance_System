from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\HomePage_D\AI_Surveillance_System")
REFERENCE = ROOT / "docs" / "NguyenMinhTuan_E11_BCDK1.docx"
OUTPUT = ROOT / "docs" / "NguyenMinhTuan_E11_BCDK2.docx"
WORK = ROOT / ".codex_doc_review" / "bcdk2"


def copy_rpr(paragraph):
    for run in paragraph.runs:
        if run._element.rPr is not None:
            return deepcopy(run._element.rPr)
    return None


def replace_text(paragraph, text):
    rpr = copy_rpr(paragraph)
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._element.insert(0, rpr)
    return paragraph


def replace_exact(document, old, new):
    matches = [p for p in document.paragraphs if p.text.strip() == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one match for {old!r}, found {len(matches)}")
    replace_text(matches[0], new)


def set_cell_text(cell, text, *, bold=False, color=None, align=None):
    paragraph = cell.paragraphs[0]
    rpr = copy_rpr(paragraph)
    for extra in list(cell.paragraphs[1:]):
        cell._element.remove(extra._element)
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    run = paragraph.add_run(str(text))
    if rpr is not None:
        run._element.insert(0, rpr)
    set_run_font(run, 10, bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    if align is not None:
        paragraph.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=13, *, bold=None, color=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        shade_cell(table.rows[0].cells[index], "2F5496")
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            color=(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                shade_cell(cells[index], "EAF0F8")
            set_cell_text(cells[index], value)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_body(document, text, *, bold_lead=None):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, 13, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, 13)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 13)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="Normal")
        paragraph_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "25")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        paragraph_pr.append(num_pr)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(item)
        set_run_font(run, 12.5)


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_font(
            run,
            {1: 20, 2: 16, 3: 14, 4: 13}.get(level, 13),
            bold=True,
            color=(47, 84, 150),
        )
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, 11, italic=True)


def font(path, size):
    return ImageFont.truetype(str(path), size)


def make_diagram(path, title, boxes, arrows):
    image = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(image)
    regular = Path(r"C:\Windows\Fonts\arial.ttf")
    bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    draw.text((750, 55), title, font=font(bold, 42), fill="#1F3864", anchor="mm")
    for box in boxes:
        x1, y1, x2, y2, heading, lines, fill = box
        draw.rounded_rectangle(
            (x1, y1, x2, y2),
            radius=22,
            fill=fill,
            outline="#2F5496",
            width=4,
        )
        draw.text(
            ((x1 + x2) / 2, y1 + 38),
            heading,
            font=font(bold, 28),
            fill="#1F3864",
            anchor="mm",
        )
        y = y1 + 82
        for line in lines:
            draw.text(
                ((x1 + x2) / 2, y),
                line,
                font=font(regular, 22),
                fill="#222222",
                anchor="mm",
            )
            y += 32
    for start, end, label in arrows:
        draw.line((start, end), fill="#2F5496", width=7)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            points = [(ex, ey), (ex - 20, ey - 12), (ex - 20, ey + 12)]
        elif ex < sx:
            points = [(ex, ey), (ex + 20, ey - 12), (ex + 20, ey + 12)]
        elif ey > sy:
            points = [(ex, ey), (ex - 12, ey - 20), (ex + 12, ey - 20)]
        else:
            points = [(ex, ey), (ex - 12, ey + 20), (ex + 12, ey + 20)]
        draw.polygon(points, fill="#2F5496")
        if label:
            label_x = (sx + ex) / 2
            label_y = (sy + ey) / 2 - 18
            label_font = font(regular, 19)
            bounds = draw.textbbox(
                (label_x, label_y),
                label,
                font=label_font,
                anchor="mm",
            )
            draw.rounded_rectangle(
                (
                    bounds[0] - 8,
                    bounds[1] - 4,
                    bounds[2] + 8,
                    bounds[3] + 4,
                ),
                radius=7,
                fill="white",
            )
            draw.text(
                (label_x, label_y),
                label,
                font=label_font,
                fill="#2F5496",
                anchor="mm",
            )
    image.save(path)


def build_diagrams():
    WORK.mkdir(parents=True, exist_ok=True)
    architecture = WORK / "architecture.png"
    make_diagram(
        architecture,
        "Production architecture",
        [
            (70, 250, 375, 560, "Flutter client", ["Dashboard", "Members / Register", "Surveillance", "Logs / Alerts"], "#EAF0F8"),
            (455, 190, 805, 620, "FastAPI backend", ["REST API", "Job managers", "Persistence services", "Readiness"], "#DDEBF7"),
            (885, 130, 1245, 380, "Shared AI engine", ["YOLOv8 + ByteTrack", "InsightFace gallery", "Recognition buffer"], "#E2F0D9"),
            (885, 485, 1245, 710, "SQLite + files", ["Members / cameras", "Live sessions / logs", "Alerts / snapshots"], "#FFF2CC"),
        ],
        [
            ((375, 405), (455, 405), "HTTP"),
            ((805, 300), (885, 260), "frames"),
            ((805, 510), (885, 585), "live only"),
        ],
    )
    modes = WORK / "persistence_modes.png"
    make_diagram(
        modes,
        "Two processing modes and their persistence boundary",
        [
            (80, 180, 430, 390, "Live webcam", ["Camera source", "Continuous worker", "Persistent session"], "#DDEBF7"),
            (80, 490, 430, 700, "Uploaded video", ["Temporary file", "Background job", "TTL results"], "#FCE4D6"),
            (575, 265, 940, 605, "Shared analysis", ["Person tracking", "Face recognition", "Known / unknown", "Annotated frame"], "#E2F0D9"),
            (1080, 180, 1430, 390, "Persistent history", ["Detection logs", "Unknown alerts", "Snapshots"], "#FFF2CC"),
            (1080, 490, 1430, 700, "Temporary output", ["Progress", "Events in memory", "Deleted upload"], "#F4CCCC"),
        ],
        [
            ((430, 285), (575, 350), ""),
            ((430, 595), (575, 520), ""),
            ((940, 350), (1080, 285), "save"),
            ((940, 520), (1080, 595), "do not save"),
        ],
    )
    return architecture, modes


def restore_preserve_only_parts():
    preserve_parts = (
        "word/header1.xml",
        "word/footer1.xml",
        "word/media/image1.png",
        "word/media/image2.png",
    )
    with ZipFile(REFERENCE, "r") as source, ZipFile(OUTPUT, "r") as generated:
        with NamedTemporaryFile(
            suffix=".docx",
            dir=OUTPUT.parent,
            delete=False,
        ) as temporary_file:
            temporary_path = Path(temporary_file.name)
        with ZipFile(temporary_path, "w", ZIP_DEFLATED) as rewritten:
            for item in generated.infolist():
                content = (
                    source.read(item.filename)
                    if item.filename in preserve_parts
                    else generated.read(item.filename)
                )
                rewritten.writestr(item, content)
    try:
        temporary_path.replace(OUTPUT)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def update_front_matter(document):
    replace_exact(document, "BÁO CÁO ĐỊNH KỲ", "BÁO CÁO ĐỊNH KỲ LẦN 2")
    replace_exact(
        document,
        "Thành phố Hồ Chí Minh, tháng 4 năm 2025",
        "Thành phố Hồ Chí Minh, tháng 7 năm 2026",
    )
    replace_exact(
        document,
        "Trong giai đoạn đầu, nhóm đã tiến hành thu thập và tiền xử lý dữ liệu hình ảnh; nghiên cứu, thử nghiệm và đánh giá các mô hình AI thông qua quy trình Jupyter Notebook; đồng thời xây dựng cấu trúc cơ sở dữ liệu và giao diện tổng thể của ứng dụng. Pipeline thử nghiệm hiện đã thực hiện được các bước phát hiện người, phát hiện và nhận dạng khuôn mặt, theo dõi đối tượng và thử nghiệm một số quy tắc cảnh báo dựa trên dữ liệu video.",
        "Trong giai đoạn báo cáo lần 2, nhóm chuyển trọng tâm từ pipeline trình diễn sang kiến trúc sản phẩm. Backend được tổ chức lại thành các module cấu hình, dữ liệu, AI, camera, dịch vụ và API; cơ sở dữ liệu được mở rộng cho camera, phiên giám sát, nhật ký, cảnh báo và vùng giám sát; đồng thời pipeline YOLOv8, ByteTrack và InsightFace được đóng gói thành một bộ máy phân tích khung hình dùng chung.",
    )
    replace_exact(
        document,
        "Bên cạnh quá trình thử nghiệm mô hình, nhóm đã chuyển đổi một phần pipeline AI từ Jupyter Notebook thành các module Python có thể sử dụng trong Backend FastAPI. Chức năng đăng ký khuôn mặt đã được hoàn thiện theo luồng end-to-end: người dùng nhập thông tin và lựa chọn hình ảnh trên ứng dụng Flutter; Backend tiếp nhận, kiểm tra và xử lý ảnh; phát hiện khuôn mặt, trích xuất vector đặc trưng bằng InsightFace; sau đó lưu thông tin người dùng, ảnh khuôn mặt và embedding vào hệ thống. Ứng dụng hiện cũng hỗ trợ hiển thị danh sách thành viên đã đăng ký và xóa thông tin thành viên.",
        "Hai chế độ xử lý đã được tích hợp end-to-end. Chế độ giám sát webcam tạo phiên làm việc, hiển thị khung hình chú thích và chỉ lưu các sự kiện có ý nghĩa, gồm nhật ký người đã biết và cảnh báo người lạ đã được xác nhận. Chế độ tải video lên sử dụng cùng bộ máy AI nhưng chỉ giữ tiến độ, khung hình và sự kiện trong bộ nhớ; tệp tải lên được xóa sau khi xử lý và không tạo dữ liệu giám sát trong SQLite.",
    )
    replace_exact(
        document,
        "Tuy nhiên, hệ thống hiện mới hoàn thành chức năng đăng ký và quản lý thành viên, cùng với các thử nghiệm AI độc lập. Những chức năng như nhận dạng thời gian thực từ camera, so khớp người quen và người lạ, tích hợp ByteTrack vào Backend, lưu nhật ký giám sát, tạo cảnh báo và kết nối đầy đủ với các màn hình Surveillance và Logs vẫn đang trong quá trình phát triển. Đây sẽ là những nội dung chính được nhóm tiếp tục thực hiện trong giai đoạn tiếp theo.",
        "Ứng dụng Flutter hiện đã kết nối với các API camera, giám sát, video tạm thời, nhật ký và cảnh báo. Dashboard hiển thị trạng thái Backend, camera, số thành viên và cảnh báo chưa đọc; màn hình Surveillance tách rõ hai chế độ; màn hình Logs chỉ hiển thị dữ liệu bền vững của webcam. Các chức năng ROI, luật hành vi nâng cao, xác thực tài khoản và thông báo đẩy chưa thuộc phạm vi triển khai của giai đoạn này.",
    )
    replace_exact(
        document,
        "Báo cáo định kỳ lần 1 trình bày tổng quan đề tài, cơ sở lý thuyết và công nghệ được sử dụng, kết quả nghiên cứu và triển khai trong giai đoạn đầu, mức độ hoàn thành của từng hạng mục, những hạn chế còn tồn tại và kế hoạch phát triển tiếp theo. Qua đó, báo cáo giúp đánh giá khách quan tiến độ hiện tại và làm cơ sở để nhóm tiếp tục hoàn thiện hệ thống trong các giai đoạn sau.",
        "Báo cáo định kỳ lần 2 kế thừa phần tổng quan, cơ sở lý thuyết và phân tích yêu cầu của báo cáo lần 1; đồng thời bổ sung thiết kế triển khai, cơ chế phân tách dữ liệu giữa hai chế độ, kết quả tích hợp Flutter–FastAPI, kiểm thử tự động và kiểm thử smoke. Báo cáo cũng nêu rõ giới hạn hiện tại để tránh đồng nhất pipeline trình diễn với mã nguồn sản phẩm.",
    )

    replace_exact(
        document,
        "1.3.4 Phạm vi của báo cáo định kỳ lần 1",
        "1.3.4 Phạm vi của báo cáo định kỳ lần 2",
    )
    replace_exact(
        document,
        "Tại thời điểm thực hiện báo cáo định kỳ lần 1, nhóm đã đạt được các kết quả chính sau:",
        "Tại thời điểm thực hiện báo cáo định kỳ lần 2, nhóm đã đạt được các kết quả chính sau:",
    )
    replacements = {
        "Hoàn thành bước thu thập và tiền xử lý dữ liệu ban đầu.": "Ổn định cấu trúc Backend FastAPI, cấu hình đường dẫn và cơ chế khởi tạo tài nguyên.",
        "Xây dựng và chạy thử pipeline AI trong Jupyter Notebook.": "Xây dựng bộ máy AI dùng chung gồm phát hiện người, ByteTrack, nhận dạng khuôn mặt và bộ đệm theo track.",
        "Thử nghiệm nhận dạng khuôn mặt, YOLOv8, ByteTrack và một số quy tắc cảnh báo.": "Hoàn thiện quản lý camera và vòng đời giám sát trực tiếp từ webcam.",
        "Xây dựng giao diện tổng thể của ứng dụng Flutter.": "Kết nối đầy đủ giao diện Flutter với Backend cho Dashboard, Surveillance, Logs và Alerts.",
        "Xây dựng cơ sở dữ liệu cho chức năng đăng ký khuôn mặt.": "Mở rộng SQLite lên schema phiên bản 2 với camera, phiên giám sát, nhật ký, cảnh báo và vùng.",
        "Xây dựng các API đăng ký, xem danh sách, xem chi tiết và xóa thành viên.": "Xây dựng các API camera, giám sát, nhật ký, cảnh báo và phân tích video tạm thời, đồng thời giữ nguyên API thành viên.",
        "Chuyển một phần pipeline AI thành các module Python trong Backend.": "Thiết lập ranh giới lưu trữ: webcam lưu lịch sử; video tải lên không ghi phiên, nhật ký, cảnh báo hoặc snapshot.",
        "Hoàn thiện chức năng đăng ký khuôn mặt và quản lý thành viên theo luồng end-to-end.": "Bổ sung readiness, script khởi động, smoke test và tài liệu vận hành cho sản phẩm.",
        "Các chức năng nhận dạng từ webcam, giám sát thời gian thực, lưu nhật ký và tạo cảnh báo chưa được tích hợp hoàn chỉnh vào sản phẩm.": "Kết quả kiểm thử hiện đạt 87 kiểm thử Backend và 4 kiểm thử Flutter; tuy nhiên chưa thực hiện benchmark mô hình thật trên bộ dữ liệu đánh giá chính thức và chưa triển khai ROI, xác thực hoặc push notification.",
    }
    for old, new in replacements.items():
        replace_exact(document, old, new)

    progress = document.tables[1]
    updates = {
        3: ("100%", "01/07–28/07"),
        4: ("100%", "05/07–28/07"),
        5: ("100%", "09/07–28/07"),
        6: ("90%", "12/07–28/07"),
    }
    for row_index, (percent, dates) in updates.items():
        set_cell_text(progress.rows[row_index].cells[3], dates)
        set_cell_text(
            progress.rows[row_index].cells[4],
            percent,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    new_rows = [
        (
            "7",
            "Xây dựng quản lý camera, phiên webcam, nhật ký và cảnh báo bền vững",
            "Huỳnh Bá Anh Khoa, Bùi Mạnh Khôi",
            "15/07–25/07",
            "95%",
        ),
        (
            "8",
            "Xây dựng chế độ phân tích video tải lên không lưu cơ sở dữ liệu",
            "Trần Tuấn Hải, Huỳnh Bá Anh Khoa",
            "20/07–27/07",
            "95%",
        ),
        (
            "9",
            "Kiểm thử tự động, readiness, smoke test và tài liệu vận hành",
            "Huỳnh Bá Anh Khoa",
            "24/07–28/07",
            "100%",
        ),
    ]
    for values in new_rows:
        cells = progress.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(
                cells[index],
                value,
                bold=index == 4,
                align=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index in (0, 3, 4)
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )


def append_report(document, architecture, modes):
    document.add_page_break()
    add_heading(document, "3.4 Đối chiếu yêu cầu với phạm vi triển khai lần 2", 2)
    add_body(
        document,
        "Các Use Case ở Chương 3 mô tả kiến trúc mục tiêu của toàn bộ đề tài. Trong báo cáo lần 2, nhóm đã hiện thực các luồng cốt lõi UC-01, UC-02, UC-05, UC-06, UC-07, UC-08 và một phần UC-09, UC-11. Cảnh báo hiện tập trung vào người lạ được xác nhận qua nhiều khung hình và cơ chế chống trùng lặp theo thời gian.",
    )
    add_body(
        document,
        "UC-03, UC-04, UC-10 và UC-12 vẫn là hạng mục kế hoạch. Cơ sở dữ liệu đã có bảng zones để chuẩn bị cho vùng giám sát, nhưng giao diện vẽ ROI và dịch vụ luật hành vi chưa được triển khai. Tương tự, ứng dụng chưa có xác thực tài khoản, quản lý thiết bị hoặc Firebase Cloud Messaging. Việc đối chiếu này giúp phân biệt rõ yêu cầu mục tiêu với chức năng đã được kiểm chứng trong mã nguồn.",
    )

    document.add_page_break()
    add_heading(document, "CHƯƠNG 4: THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG", 1)
    add_heading(document, "4.1 Kiến trúc mã nguồn sản phẩm", 2)
    add_body(
        document,
        "Mã nguồn sản phẩm được tách khỏi thư mục beta_testing_files_02. Thư mục thử nghiệm của nhóm được dùng để tham khảo pipeline và trình diễn, trong khi Backend vận hành thực tế nằm trong backend/ và ứng dụng người dùng nằm trong app_flutter/. Cách tổ chức này tránh phụ thuộc trực tiếp vào Notebook hoặc script trình diễn và cho phép kiểm thử từng lớp độc lập.",
    )
    document.add_picture(str(architecture), width=Inches(6.45))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, "Hình 4.1. Kiến trúc triển khai của mã nguồn sản phẩm")
    add_table(
        document,
        ["Lớp", "Thành phần chính", "Trách nhiệm"],
        [
            ("Flutter", "Dashboard, Members, Register, Surveillance, Logs", "Hiển thị trạng thái, gửi yêu cầu và phân tách rõ hai chế độ xử lý."),
            ("API", "FastAPI routes và schema Pydantic", "Kiểm tra dữ liệu, trả lỗi HTTP nhất quán và giữ hợp đồng API thành viên."),
            ("Dịch vụ", "Camera, live manager, video manager, log, alert", "Điều phối vòng đời công việc và quyết định dữ liệu nào được lưu."),
            ("AI", "YOLOv8, ByteTrack, InsightFace, recognition buffer", "Phân tích từng khung hình, không tự ghi cơ sở dữ liệu."),
            ("Dữ liệu", "SQLite và thư mục runtime", "Lưu thành viên và lịch sử webcam; quản lý snapshot cảnh báo."),
        ],
        [1150, 2900, 5310],
    )

    add_heading(document, "4.2 Ổn định Backend và quản lý cấu hình", 2)
    add_body(
        document,
        "Backend sử dụng một entry point duy nhất tại backend.main. FastAPI lifespan tạo thư mục runtime, khởi tạo schema SQLite khi ứng dụng bắt đầu và yêu cầu hai worker giải phóng tài nguyên khi ứng dụng dừng. Các mô hình AI được khởi tạo lười, vì vậy các API chỉ đọc như health, members hoặc logs không tải InsightFace và YOLO không cần thiết.",
    )
    add_body(
        document,
        "Các đường dẫn cơ sở dữ liệu, ảnh khuôn mặt, embedding, snapshot, model và tệp video tạm thời được tập trung trong app/config.py. Đường dẫn tương đối luôn được giải quyết từ thư mục backend thay vì phụ thuộc thư mục hiện hành. Các tham số ngưỡng nhận dạng, số khung hình xác nhận người lạ, thời gian cooldown, kích thước video tối đa và TTL kết quả đều có thể cấu hình bằng biến môi trường.",
    )
    add_heading(document, "4.3 Cơ sở dữ liệu và ranh giới lưu trữ", 2)
    add_body(
        document,
        "Schema SQLite phiên bản 2 gồm people, face_embeddings, cameras, surveillance_sessions, detection_logs, alerts và zones. Khóa ngoại và chỉ mục được thiết lập để giữ tính toàn vẹn và hỗ trợ truy vấn theo camera, phiên hoặc thời gian. Khi xóa thành viên, lịch sử giám sát vẫn được giữ nhưng liên kết member_id được đặt NULL; khi xóa phiên, logs và alerts thuộc phiên được xóa theo cascade.",
    )
    document.add_picture(str(modes), width=Inches(6.45))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, "Hình 4.2. Ranh giới lưu trữ giữa webcam và video tải lên")
    add_table(
        document,
        ["Dữ liệu", "Webcam trực tiếp", "Video tải lên"],
        [
            ("Phiên xử lý", "Lưu surveillance_sessions", "Không có bảng job hoặc session"),
            ("Sự kiện nhận dạng", "Lưu sự kiện có ý nghĩa", "Giữ tạm trong RAM"),
            ("Cảnh báo người lạ", "Lưu alerts sau xác nhận", "Chỉ hiển thị tạm thời"),
            ("Ảnh minh chứng", "Lưu snapshot cho cảnh báo", "Không lưu snapshot"),
            ("Tệp đầu vào", "Nguồn camera, không sao chép", "Xóa sau hoàn thành/lỗi/hủy"),
            ("Thời gian tồn tại", "Theo vòng đời dữ liệu SQLite", "TTL mặc định 30 phút"),
        ],
        [2400, 3320, 3640],
    )

    add_heading(document, "4.4 Bộ máy AI dùng chung", 2)
    add_body(
        document,
        "FrameAnalysisEngine là lõi không lưu trạng thái bền vững. Mỗi khung hình được chuyển qua bộ theo dõi người YOLOv8 kết hợp ByteTrack. Với từng track, hệ thống giới hạn bounding box vào kích thước ảnh, trích vùng quan tâm và gọi bộ nhận dạng khuôn mặt. Gallery trong bộ nhớ so sánh cosine similarity giữa embedding mới với embedding của các thành viên.",
    )
    add_body(
        document,
        "TrackRecognitionBuffer duy trì một cửa sổ kết quả riêng cho từng track. Cơ chế này giảm dao động danh tính giữa các khung hình và cho phép track kế thừa danh tính khi khuôn mặt tạm thời bị khuất. Trạng thái low_quality được phân biệt với unknown để khung hình không có khuôn mặt rõ không tạo cảnh báo sai. Engine trả về kiểu dữ liệu có cấu trúc gồm track_id, bounding box, độ tin cậy người, trạng thái, member_id và similarity.",
    )
    add_heading(document, "4.5 Quản lý camera và giám sát webcam", 2)
    add_body(
        document,
        "API camera hỗ trợ tạo, liệt kê, cập nhật, vô hiệu hóa mềm, kiểm tra nguồn và lấy ảnh xem trước. Nguồn chuỗi số như \"0\" được chuyển thành chỉ số webcam OpenCV. Mỗi lần kiểm tra hoặc chụp preview đều mở nguồn trong thời gian ngắn và luôn release tài nguyên. Preview bị từ chối nếu camera đang thuộc quyền sở hữu của worker giám sát.",
    )
    add_body(
        document,
        "Khi bắt đầu giám sát, Backend kiểm tra camera tồn tại và đang hoạt động, tạo surveillance_session, làm mới gallery và chạy worker nền. Worker cập nhật số khung hình, FPS, kết quả mới nhất và JPEG chú thích trong bộ nhớ để Flutter polling. Khi dừng hoặc gặp lỗi, camera được giải phóng và phiên được cập nhật stopped hoặc failed cùng số khung hình, FPS và thông báo lỗi.",
    )
    add_heading(document, "4.6 Nhật ký và cảnh báo webcam", 2)
    add_body(
        document,
        "LiveEventRecorder không ghi một bản ghi cho mọi khung hình. Với người đã biết, hệ thống chỉ ghi khi track lần đầu được nhận dạng hoặc danh tính thay đổi. Với người lạ, track phải duy trì trạng thái unknown qua số khung hình cấu hình trước khi tạo detection_log và alert. Cơ chế cooldown hạn chế nhiều cảnh báo liên tiếp của cùng track.",
    )
    add_body(
        document,
        "Snapshot cảnh báo chỉ được tạo cho sự kiện người lạ đã xác nhận. Nếu giao dịch cơ sở dữ liệu thất bại, file snapshot mới được xóa để tránh dữ liệu mồ côi. API logs hỗ trợ lọc theo trạng thái, camera và phiên; API alerts hỗ trợ danh sách, chi tiết, latest, ảnh minh chứng và cập nhật trạng thái đã đọc.",
    )
    add_heading(document, "4.7 Phân tích video tải lên tạm thời", 2)
    add_body(
        document,
        "VideoAnalysisManager tiếp nhận tệp đa phần, xác thực phần mở rộng và kích thước, sau đó chạy một job nền. Trạng thái queued, running, completed, failed hoặc cancelled cùng tiến độ, FPS, khung hình mới nhất và sự kiện được giữ trong bộ nhớ. Danh sách sự kiện có giới hạn để tránh tăng bộ nhớ không kiểm soát.",
    )
    add_body(
        document,
        "TemporaryEventCollector áp dụng nguyên tắc xác nhận unknown và cooldown tương tự live mode nhưng không gọi LogService, AlertService hoặc SQLite. Tệp video được xóa trong mọi đường kết thúc. Chỉ một job video được phép hoạt động và job video loại trừ lẫn nhau với live surveillance, giúp giảm tranh chấp GPU/CPU trong môi trường trình diễn.",
    )
    add_heading(document, "4.8 Tích hợp ứng dụng Flutter", 2)
    add_body(
        document,
        "Flutter sử dụng security_models.dart để ánh xạ response typed và security_service.dart để tập trung HTTP, multipart upload, xử lý lỗi và tải JPEG. Màn hình Surveillance có hai lựa chọn rõ ràng. Live webcam cho phép thêm nguồn laptop, chọn camera, bắt đầu/dừng, xem FPS, khung hình và track. Upload video hiển thị tiến độ, sự kiện tạm thời, hủy job và xóa kết quả.",
    )
    add_body(
        document,
        "Dashboard polling tình trạng Backend, camera, số thành viên, số cảnh báo chưa đọc và detection mới nhất. Màn hình Persistent Logs & Alerts chỉ gọi API dữ liệu webcam và cho phép đánh dấu cảnh báo đã đọc/chưa đọc. Nội dung hướng dẫn trên giao diện nhấn mạnh rằng kết quả video tải lên không xuất hiện trong lịch sử.",
    )

    document.add_page_break()
    add_heading(document, "CHƯƠNG 5: KẾT QUẢ TRIỂN KHAI VÀ KIỂM THỬ", 1)
    add_heading(document, "5.1 Chiến lược kiểm thử", 2)
    add_body(
        document,
        "Kiểm thử được tổ chức theo lớp. Unit test kiểm tra hợp đồng dữ liệu, buffer nhận dạng, gallery, engine phân tích, collector sự kiện và dịch vụ lưu trữ. API test sử dụng FastAPI TestClient cùng connection SQLite trong bộ nhớ. Các model nặng được thay bằng test double để kiểm thử không tải model hoặc phụ thuộc webcam thật. Flutter test kiểm tra ánh xạ JSON, lỗi Backend và khả năng hiển thị shell ứng dụng.",
    )
    add_body(
        document,
        "Ngoài kiểm thử tự động, readiness endpoint kiểm tra schema, thư mục ghi dữ liệu và trạng thái worker mà không tải AI. Smoke test khởi động Backend tạm thời trên cổng riêng và gọi bảy endpoint chỉ đọc. Quy trình này xác nhận ứng dụng thực sự khởi động và có thể phục vụ Flutter, đồng thời không tạo camera, phiên hoặc dữ liệu mới.",
    )
    add_heading(document, "5.2 Kết quả kiểm thử tự động", 2)
    add_table(
        document,
        ["Hạng mục", "Kết quả", "Bằng chứng"],
        [
            ("Backend unit/API", "87/87 đạt", "unittest discover; bao phủ AI contract, DB, camera, live, logs, alerts, upload và readiness"),
            ("Flutter unit/widget", "4/4 đạt", "Ánh xạ camera/status, non-persistent marker, lỗi API và shell 5 màn hình"),
            ("Flutter static analysis", "Không có lỗi", "flutter analyze: No issues found"),
            ("OpenAPI", "26 paths", "Health, members, cameras, surveillance, logs, alerts và video-analysis"),
            ("Smoke test", "7/7 endpoint đạt", "health, readiness, members, cameras, surveillance/status, logs, alerts"),
            ("Repository hygiene", "Đạt", "beta_testing_files_02 không thay đổi; file sinh tự động được khôi phục"),
        ],
        [2500, 1800, 5060],
    )
    add_heading(document, "5.3 Các kịch bản quan trọng đã kiểm chứng", 2)
    add_bullets(
        document,
        [
            "Camera không tồn tại hoặc bị vô hiệu hóa không thể bắt đầu giám sát.",
            "Camera được release khi đọc ảnh thành công, thất bại hoặc worker kết thúc.",
            "Không thể chạy live surveillance đồng thời với job video tải lên.",
            "Người lạ phải được xác nhận qua nhiều khung hình và tuân thủ cooldown.",
            "Khung hình low_quality không làm tăng chuỗi xác nhận người lạ.",
            "Lỗi giao dịch xóa snapshot vừa tạo, không để lại file mồ côi.",
            "Job video hoàn thành, lỗi hoặc hủy đều xóa tệp đầu vào tạm.",
            "Phân tích video tải lên không làm thay đổi surveillance_sessions, detection_logs hoặc alerts.",
            "Xóa thành viên giữ lịch sử giám sát và xóa liên kết danh tính theo đúng khóa ngoại.",
            "Readiness phát hiện schema thiếu bảng và trả trạng thái degraded.",
        ],
    )
    add_heading(document, "5.4 Kết quả chức năng theo hai chế độ", 2)
    add_table(
        document,
        ["Tiêu chí", "Live webcam", "Uploaded video"],
        [
            ("Bắt đầu từ Flutter", "Chọn camera hoặc thêm laptop webcam", "Chọn tệp video hỗ trợ"),
            ("Theo dõi trạng thái", "running, FPS, frames, lỗi", "queued/running/completed, progress, FPS"),
            ("Khung hình", "JPEG chú thích mới nhất", "JPEG chú thích mới nhất"),
            ("Kết quả", "Track nhận dạng hiện tại", "Danh sách sự kiện tạm thời"),
            ("Lịch sử", "Có logs và alerts", "Không lưu lịch sử"),
            ("Kết thúc", "Cập nhật phiên và release camera", "Xóa upload; kết quả hết hạn theo TTL"),
        ],
        [2300, 3530, 3530],
    )
    add_heading(document, "5.5 Khả năng vận hành và chẩn đoán", 2)
    add_body(
        document,
        "Nhóm bổ sung bốn script: start_backend.ps1 kiểm tra dependency và chạy Uvicorn; start_flutter.ps1 yêu cầu readiness ở trạng thái ready trước khi mở ứng dụng; smoke_test.py thực hiện kiểm tra chỉ đọc; verify_project.ps1 chạy toàn bộ test Backend, flutter analyze và flutter test. README ở thư mục gốc được cập nhật thành runbook thống nhất thay cho mô tả chỉ có chức năng đăng ký.",
    )
    add_body(
        document,
        "Endpoint /health/readiness trả trạng thái database, storage, AI assets và worker. Trạng thái download_on_first_use của model chỉ mang tính thông tin và không làm Backend degraded, vì hệ thống chủ động lazy load. Backend chỉ degraded khi schema không đúng hoặc thư mục runtime không sẵn sàng ghi.",
    )

    document.add_page_break()
    add_heading(document, "CHƯƠNG 6: ĐÁNH GIÁ VÀ KẾ HOẠCH TIẾP THEO", 1)
    add_heading(document, "6.1 Mức độ hoàn thành theo phân công", 2)
    add_body(
        document,
        "Phần AI đã cung cấp pipeline phát hiện, theo dõi và nhận dạng có thể gọi lại từ Backend. Phần dữ liệu đã hình thành schema và ràng buộc cho các thực thể cốt lõi. Phần ứng dụng và tích hợp đã chuyển pipeline thành module sản phẩm, xây dựng API, hai worker, giao diện Flutter và bộ kiểm thử. Các phần giao nhau được xác nhận bằng test contract thay vì gắn trực tiếp mã Notebook vào ứng dụng.",
    )
    add_body(
        document,
        "So với báo cáo lần 1, hạng mục tích hợp AI Backend và Flutter tăng từ 30% lên khoảng 90%. Phần còn lại chủ yếu là đánh giá mô hình trên dữ liệu thật, tối ưu hiệu năng phần cứng, hoàn thiện ROI/behavior rules và kiểm thử trình diễn với webcam/video đại diện.",
    )
    add_heading(document, "6.2 Hạn chế hiện tại", 2)
    add_bullets(
        document,
        [
            "Bộ kiểm thử tự động dùng test double và chưa benchmark YOLOv8/InsightFace thật trong cùng môi trường triển khai.",
            "Chưa có số liệu Accuracy, FAR, FRR, Precision, Recall, F1-score và FPS chính thức trên bộ dữ liệu đánh giá chung.",
            "Chưa có giao diện vẽ ROI, API CRUD zones hoặc dịch vụ behavior rules.",
            "Cảnh báo mới tập trung vào unknown_person; restricted_area và loitering chưa triển khai.",
            "Chưa có đăng nhập, phân quyền, quản lý thiết bị hoặc Firebase push notification.",
            "Worker hiện giới hạn một pipeline nặng tại một thời điểm và chưa hướng đến nhiều camera đồng thời.",
            "SQLite và lưu trữ cục bộ phù hợp trình diễn nhưng chưa phải kiến trúc triển khai phân tán.",
        ],
    )
    add_heading(document, "6.3 Kế hoạch giai đoạn tiếp theo", 2)
    add_table(
        document,
        ["Ưu tiên", "Hạng mục", "Tiêu chí hoàn thành"],
        [
            ("1", "Kiểm thử thật end-to-end", "Chạy webcam và video mẫu với model thật; ghi FPS, latency và lỗi."),
            ("2", "Đánh giá nhận dạng", "Xây dựng tập gallery/query và báo cáo FAR, FRR, Precision, Recall, F1."),
            ("3", "ROI và luật hành vi", "CRUD zones, vẽ polygon Flutter, restricted_area và loitering."),
            ("4", "Bảo mật", "Xác thực, phân quyền API, bảo vệ dữ liệu khuôn mặt và audit log."),
            ("5", "Thông báo", "FCM hoặc kênh thông báo phù hợp cho alert mới."),
            ("6", "Hoàn thiện báo cáo", "Bổ sung ảnh chạy thật, biểu đồ hiệu năng và kết quả đánh giá."),
        ],
        [1000, 3000, 5360],
    )

    document.add_page_break()
    add_heading(document, "KẾT LUẬN", 1)
    add_body(
        document,
        "Trong giai đoạn báo cáo định kỳ lần 2, nhóm đã chuyển hệ thống từ trạng thái đăng ký khuôn mặt và pipeline thử nghiệm sang một sản phẩm tích hợp có kiến trúc rõ ràng. Backend FastAPI quản lý camera, phiên webcam, engine AI dùng chung, nhật ký, cảnh báo và job video tạm thời. Flutter đã kết nối với các API này để vận hành hai chế độ và hiển thị đúng ranh giới dữ liệu.",
    )
    add_body(
        document,
        "Điểm quan trọng của thiết kế là nguyên tắc lưu trữ theo ngữ cảnh: webcam phục vụ giám sát thực tế nên lưu phiên và sự kiện có ý nghĩa; video tải lên phục vụ phân tích/trình diễn nên không tạo lịch sử. Cơ chế xác nhận người lạ, cooldown, cleanup file và mutual exclusion giúp giảm cảnh báo trùng, dữ liệu mồ côi và tranh chấp tài nguyên.",
    )
    add_body(
        document,
        "Kết quả 87 kiểm thử Backend, 4 kiểm thử Flutter, static analysis sạch và smoke test thực tế cho thấy các hợp đồng phần mềm chính đã ổn định. Tuy nhiên, kết quả này chưa thay thế benchmark mô hình thật. Giai đoạn tiếp theo cần tập trung vào đánh giá định lượng, thử nghiệm webcam/video thực, ROI, luật hành vi và bảo mật trước khi xem hệ thống là hoàn chỉnh.",
    )

    add_heading(document, "TÀI LIỆU THAM KHẢO", 1)
    references = [
        "[1] FastAPI Documentation, https://fastapi.tiangolo.com/.",
        "[2] Flutter Documentation, https://docs.flutter.dev/.",
        "[3] Ultralytics YOLO Documentation, https://docs.ultralytics.com/.",
        "[4] Y. Zhang et al., “ByteTrack: Multi-Object Tracking by Associating Every Detection Box,” ECCV, 2022.",
        "[5] J. Deng et al., “ArcFace: Additive Angular Margin Loss for Deep Face Recognition,” CVPR, 2019.",
        "[6] InsightFace Project, https://github.com/deepinsight/insightface.",
        "[7] SQLite Documentation, https://www.sqlite.org/docs.html.",
        "[8] OpenCV Documentation, https://docs.opencv.org/.",
    ]
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(reference)
        set_run_font(run, 12)

    document.add_page_break()
    add_heading(document, "PHỤ LỤC A: DANH SÁCH API ĐÃ TRIỂN KHAI", 1)
    api_rows = [
        ("System", "GET /health; GET /health/readiness", "Trạng thái và chẩn đoán"),
        ("Members", "POST /members/register; GET /members; GET/DELETE /members/{id}", "Đăng ký và quản lý thành viên"),
        ("Cameras", "POST/GET /cameras; GET/PATCH/DELETE /cameras/{id}", "Quản lý cấu hình camera"),
        ("Cameras", "POST /cameras/{id}/test; GET /cameras/{id}/snapshot", "Kiểm tra và preview"),
        ("Live", "POST /surveillance/start; POST /surveillance/stop", "Điều khiển webcam"),
        ("Live", "GET /surveillance/status; /latest; /frame", "Polling trạng thái và kết quả"),
        ("Logs", "GET /logs; GET /logs/latest", "Lịch sử detection webcam"),
        ("Alerts", "GET /alerts; /latest; /{id}; /{id}/snapshot", "Cảnh báo và ảnh minh chứng"),
        ("Alerts", "PATCH /alerts/{id}/read", "Đánh dấu đã đọc/chưa đọc"),
        ("Video", "POST /video-analysis", "Tạo job tạm thời"),
        ("Video", "GET /video-analysis/{id}/status; /results; /frame", "Polling video tạm thời"),
        ("Video", "DELETE /video-analysis/{id}", "Hủy hoặc xóa job"),
    ]
    add_table(
        document,
        ["Nhóm", "Endpoint", "Mục đích"],
        api_rows,
        [1200, 4930, 3230],
    )
    add_body(
        document,
        "OpenAPI của phiên bản báo cáo lần 2 gồm 26 path. Danh sách trên gộp các phương thức có cùng tài nguyên để trình bày ngắn gọn; hợp đồng chi tiết có thể xem tại /docs khi Backend đang chạy.",
    )


def main():
    architecture, modes = build_diagrams()
    document = Document(REFERENCE)
    update_front_matter(document)
    append_report(document, architecture, modes)
    descriptions = {
        "Picture 1": "Biểu trưng của Học viện Công nghệ Bưu chính Viễn thông",
        "Picture 1983538751": "Sơ đồ Use Case tổng quan của hệ thống",
        "Picture 1983538752": "Kiến trúc triển khai Flutter, FastAPI, AI và SQLite",
        "Picture 1983538753": "Ranh giới lưu trữ giữa webcam trực tiếp và video tải lên",
    }
    for doc_pr in document.element.iter(qn("wp:docPr")):
        description = descriptions.get(
            doc_pr.get("name"),
            "Thành phần đồ họa trang trí của báo cáo",
        )
        doc_pr.set("descr", description)
        doc_pr.set("title", description)
    for table in document.tables:
        set_repeat_table_header(table.rows[0])
    core = document.core_properties
    core.title = "Báo cáo định kỳ lần 2 - AI Face Recognition Security System"
    core.subject = "Tiến độ triển khai Backend, Flutter và kiểm thử tích hợp"
    core.author = "Nhóm E11"
    core.comments = "Generated from the retained BCDK1 report structure."
    document.save(OUTPUT)
    restore_preserve_only_parts()
    print(OUTPUT)


if __name__ == "__main__":
    main()
